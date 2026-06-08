"""
Function tools definitions for AI agent conversation.
These tools are invoked by the AI agent during conversation.
"""

import asyncio
import httpx
import json
import re
import uuid
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from app.models.project import Project
from app.models.call_spec import CallSpec
from app.models.document import GeneratedDoc
from app.models.cyrano_evaluation import CyranoEvaluation
from app.models.enums import ProjectStatus, ProjectLanguage
from app.database import SessionLocal
from app.config import get_settings
from app.core.prompts import EXTRACT_REQUIREMENTS_PROMPT
from app.services.vector_store import upload_bytes_to_projects_store


# ───────────────────────────────────────────
# Tool definitions for function-calling tools
# ───────────────────────────────────────────

TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "search_funding_calls",
            "description": "Busca convocatorias activas de financiamiento en la web según sector y territorio. Retorna una lista de convocatorias relevantes.",
            "parameters": {
                "type": "object",
                "properties": {
                    "sector": {
                        "type": "string",
                        "description": "Sector del proyecto (ej: agricultura, tecnología, salud, educación)"
                    },
                    "territory": {
                        "type": "string",
                        "description": "País o región de interés"
                    },
                    "keywords": {
                        "type": "string",
                        "description": "Palabras clave adicionales para la búsqueda"
                    }
                },
                "required": ["sector", "territory", "keywords"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "extract_requirements",
            "description": "Extrae requisitos estructurados de un documento de convocatoria. Úsalo cuando el usuario suba un archivo O cuando pegue directamente el texto de una convocatoria/spec en el chat.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_text": {
                        "type": "string",
                        "description": "Texto completo del documento o convocatoria a analizar"
                    },
                    "call_spec_id": {
                        "type": "string",
                        "description": "ID de la convocatoria en la base de datos (opcional — omitir si el texto fue pegado directamente)"
                    },
                    "title": {
                        "type": "string",
                        "description": "Título descriptivo de la convocatoria"
                    }
                },
                "required": ["document_text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_budget",
            "description": "Valida que los rubros presupuestarios no superen los topes de la convocatoria.",
            "parameters": {
                "type": "object",
                "properties": {
                    "budget_items": {
                        "type": "string",
                        "description": "JSON string con los rubros del presupuesto: [{rubro, monto, actividad_vinculada}]"
                    },
                    "max_total": {
                        "type": "number",
                        "description": "Monto máximo permitido por la convocatoria"
                    },
                    "admin_cap_percent": {
                        "type": "number",
                        "description": "Porcentaje máximo permitido para gastos administrativos"
                    },
                    "project_id": {
                        "type": "string",
                        "description": "ID del proyecto para guardar el presupuesto validado"
                    }
                },
                "required": ["budget_items", "max_total", "admin_cap_percent", "project_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_word_document",
            "description": "Genera un documento Word (.docx) profesional con toda la información del proyecto.",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_id": {
                        "type": "string",
                        "description": "ID del proyecto a exportar"
                    },
                    "language": {
                        "type": "string",
                        "description": "Idioma del documento: 'es' para español, 'en' para inglés",
                        "enum": ["es", "en"]
                    },
                    "content": {
                        "type": "string",
                        "description": "Contenido completo del proyecto en formato markdown"
                    }
                },
                "required": ["project_id", "language", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_diagnostic",
            "description": "Ejecuta el diagnóstico Cyrano para evaluar la calidad del proyecto.",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_id": {
                        "type": "string",
                        "description": "ID del proyecto a diagnosticar"
                    }
                },
                "required": ["project_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_project_data",
            "description": "Guarda los datos del proyecto en PostgreSQL.",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_id": {
                        "type": "string",
                        "description": "ID del proyecto existente para actualizar."
                    },
                    "user_id": {
                        "type": "string",
                        "description": "ID del usuario propietario"
                    },
                    "title": {
                        "type": "string",
                        "description": "Título del proyecto"
                    },
                    "problem_definition": {
                        "type": "string",
                        "description": "Texto de la definición del problema"
                    },
                    "objectives": {
                        "type": "string",
                        "description": "Texto de los objetivos"
                    },
                    "methodology": {
                        "type": "string",
                        "description": "Descripción de la metodología"
                    },
                    "timeline_text": {
                        "type": "string",
                        "description": "Texto del cronograma"
                    },
                    "budget_text": {
                        "type": "string",
                        "description": "Texto del presupuesto detallado"
                    },
                    "full_content": {
                        "type": "string",
                        "description": "Contenido completo del proyecto en formato markdown"
                    }
                },
                "required": ["project_id", "user_id", "title", "problem_definition", "objectives", "methodology", "timeline_text", "budget_text", "full_content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_to_project_memory",
            "description": "Guarda un resumen del proyecto completado en el Vector Store de memoria de proyectos.",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_id": {
                        "type": "string",
                        "description": "ID del proyecto a guardar en memoria"
                    },
                    "summary": {
                        "type": "string",
                        "description": "Resumen ejecutivo del proyecto completado para memoria de conocimiento."
                    }
                },
                "required": ["project_id", "summary"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "save_diagnostic_result",
            "description": "Persiste el resultado estructurado del diagnóstico Cyrano en la base de datos.",
            "parameters": {
                "type": "object",
                "properties": {
                    "project_id": {"type": "string", "description": "ID del proyecto evaluado"},
                    "score": {"type": "number", "description": "Puntaje total ponderado (0–100)"},
                    "sections": {
                        "type": "object",
                        "description": "Puntaje de 1 a 10 por cada sección.",
                        "properties": {
                            "problem_definition": {"type": "number"},
                            "problem_tree": {"type": "number"},
                            "objectives": {"type": "number"},
                            "value_chain": {"type": "number"},
                            "timeline": {"type": "number"},
                            "budget": {"type": "number"},
                        },
                        "required": ["problem_definition", "problem_tree", "objectives", "value_chain", "timeline", "budget"],
                    },
                    "gaps": {
                        "type": "array",
                        "description": "Lista de brechas identificadas",
                        "items": {"type": "string"},
                    },
                    "recommendations": {
                        "type": "array",
                        "description": "Lista de recomendaciones prioritizadas",
                        "items": {"type": "string"},
                    },
                    "verdict": {
                        "type": "string",
                        "description": "APROBADO si score >= 95.01, EN REVISIÓN si score < 95.01",
                        "enum": ["APROBADO", "EN REVISIÓN"],
                    },
                },
                "required": ["project_id", "score", "sections", "gaps", "recommendations", "verdict"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Descarga y extrae el texto de una página web dado su URL. Úsalo cuando el usuario comparta una URL de convocatoria o sitio de financiamiento para analizar su contenido.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "URL completa de la página a descargar (ej: https://www.opentech.fund/funds/)"
                    }
                },
                "required": ["url"],
            },
        },
    },
]


# ───────────────────────────────────────────
# Tool handlers
# ───────────────────────────────────────────

async def handle_fetch_url(args: dict) -> str:
    """Fetch a web page and return its plain text content."""
    url = args.get("url", "").strip()
    if not url:
        return json.dumps({"status": "error", "message": "Se requiere una URL."})

    # Basic URL validation — must start with http/https
    if not url.startswith(("http://", "https://")):
        return json.dumps({"status": "error", "message": "URL inválida. Debe comenzar con http:// o https://"})

    try:
        from bs4 import BeautifulSoup

        headers = {
            "User-Agent": "Mozilla/5.0 (compatible; GrantsAnalytics/1.0; research bot)",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "en,es;q=0.9",
        }
        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove script, style, nav, footer noise
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)

        # Collapse excessive blank lines
        lines = [ln for ln in text.splitlines() if ln.strip()]
        clean_text = "\n".join(lines)

        # Limit to 12 000 chars to stay within context window
        MAX_CHARS = 12_000
        truncated = len(clean_text) > MAX_CHARS
        if truncated:
            clean_text = clean_text[:MAX_CHARS]

        return json.dumps({
            "status": "ok",
            "url": url,
            "content": clean_text,
            "truncated": truncated,
            "chars": len(clean_text),
        })

    except httpx.HTTPStatusError as e:
        return json.dumps({"status": "error", "url": url, "message": f"HTTP {e.response.status_code}: {e.response.reason_phrase}"})
    except Exception as e:
        return json.dumps({"status": "error", "url": url, "message": f"No se pudo obtener la página: {str(e)}"})


async def handle_search_funding_calls(args: dict) -> str:
    """Search for active funding calls using DuckDuckGo (no API key required)."""
    sector = args.get("sector", "")
    territory = args.get("territory", "")
    keywords = args.get("keywords", "")

    query_parts = []
    if keywords:
        query_parts.append(keywords)
    if sector:
        query_parts.append(sector)
    if territory:
        query_parts.append(territory)
    query_parts.append("grant funding call open 2025 2026")

    query = " ".join(query_parts)

    try:
        from ddgs import DDGS

        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=8):
                results.append({
                    "title": r.get("title", ""),
                    "url": r.get("href", ""),
                    "snippet": r.get("body", ""),
                })

        if not results:
            return json.dumps({
                "status": "ok",
                "query": query,
                "results": [],
                "message": "No se encontraron resultados para esta búsqueda.",
            })

        return json.dumps({
            "status": "ok",
            "query": query,
            "results": results,
            "message": f"Se encontraron {len(results)} convocatorias relevantes.",
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "query": query,
            "message": f"Error en la búsqueda web: {str(e)}",
        })


async def handle_extract_requirements(args: dict) -> str:
    """Extract structured requirements from a call document using Ollama."""
    document_text = args["document_text"]
    call_spec_id = args.get("call_spec_id")
    title = args.get("title", "Convocatoria")

    _settings = get_settings()
    prompt = f"{EXTRACT_REQUIREMENTS_PROMPT}\n\n<document>\n{document_text}\n</document>"

    url = _settings.OLLAMA_BASE_URL.rstrip("/") + "/v1/chat/completions"
    payload = {
        "model": _settings.OLLAMA_MODEL,
        "messages": [{"role": "user", "content": prompt}],
    }

    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(url, json=payload)
        resp.raise_for_status()
        data = resp.json()
        text = data.get("choices", [{}])[0].get("message", {}).get("content", "")

    extracted: dict = {}
    try:
        json_match = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
        if json_match:
            extracted = json.loads(json_match.group(1))
        else:
            extracted = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        extracted = {"raw_extraction": text}

    def _save_to_db() -> str:
        db = SessionLocal()
        try:
            def _str(val: object, limit: int = 250) -> str | None:
                s = str(val).strip() if val else None
                return s[:limit] if s else None

            if call_spec_id:
                spec = db.query(CallSpec).filter(
                    CallSpec.id == uuid.UUID(call_spec_id)
                ).first()
                if spec:
                    spec.extracted_requirements = extracted
                    spec.eligibility_criteria = _str(extracted.get("criterios_elegibilidad"), 2000)
                    spec.max_amount = _str(extracted.get("montos_maximos"))
                    spec.counterpart_required = _str(extracted.get("contrapartidas_requeridas"))
                    spec.deadline = _str(extracted.get("fechas_cierre"))
                    spec.mandatory_sections = extracted.get("secciones_obligatorias")
                    spec.raw_text = document_text[:10_000]
                    db.commit()
                    return str(spec.id)

            spec = CallSpec(
                title=title,
                extracted_requirements=extracted,
                eligibility_criteria=_str(extracted.get("criterios_elegibilidad"), 2000),
                max_amount=_str(extracted.get("montos_maximos")),
                counterpart_required=_str(extracted.get("contrapartidas_requeridas")),
                deadline=_str(extracted.get("fechas_cierre")),
                mandatory_sections=extracted.get("secciones_obligatorias"),
                raw_text=document_text[:10_000],
            )
            db.add(spec)
            db.commit()
            db.refresh(spec)
            return str(spec.id)
        finally:
            db.close()

    saved_id = await asyncio.to_thread(_save_to_db)

    return json.dumps({
        "status": "success",
        "call_spec_id": saved_id,
        "document_length": len(document_text),
        "extracted": extracted,
        "message": "Requisitos extraídos y guardados correctamente.",
    })


def handle_calculate_budget(args: dict, db: Session) -> str:
    """Validate budget items against call constraints."""
    try:
        budget_items = json.loads(args["budget_items"])
    except (json.JSONDecodeError, TypeError):
        return json.dumps({"status": "error", "message": "Formato de presupuesto inválido."})

    max_total = args["max_total"]
    admin_cap = args["admin_cap_percent"]
    project_id = uuid.UUID(args["project_id"]) if isinstance(args["project_id"], str) else args["project_id"]

    total = sum(item.get("monto", 0) for item in budget_items)
    admin_total = sum(item.get("monto", 0) for item in budget_items if item.get("rubro", "").lower() in ["administrativo", "admin", "gastos administrativos"])
    admin_percent = (admin_total / total * 100) if total > 0 else 0

    issues = []
    if total > max_total:
        issues.append(f"El presupuesto total ({total:,.2f}) excede el máximo permitido ({max_total:,.2f})")
    if admin_percent > admin_cap:
        issues.append(f"Los gastos administrativos ({admin_percent:.1f}%) exceden el tope ({admin_cap}%)")

    unlinked = [item for item in budget_items if not item.get("actividad_vinculada")]
    if unlinked:
        issues.append(f"{len(unlinked)} rubros sin actividad vinculada")

    result = {
        "status": "valid" if not issues else "issues_found",
        "total": total,
        "max_allowed": max_total,
        "admin_percent": round(admin_percent, 1),
        "admin_cap": admin_cap,
        "issues": issues,
        "items_count": len(budget_items),
    }

    if not issues:
        project = db.query(Project).filter(Project.id == project_id).first()
        if project:
            project.budget = {"items": budget_items, "total": total, "validated": True}
            db.commit()

    return json.dumps(result)


def handle_generate_word_document(args: dict, db: Session) -> str:
    """Generate a professional Word document for the project."""
    project_id = uuid.UUID(args["project_id"]) if isinstance(args["project_id"], str) else args["project_id"]
    language = args["language"]
    content = args.get("content")

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return json.dumps({"status": "error", "message": "Proyecto no encontrado"})

    current_score = project.cyrano_score
    is_draft = current_score is None or current_score < 95.01

    doc = Document()
    title_para = doc.add_heading(project.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if current_score is not None:
        score_label = "BORRADOR" if is_draft else "DOCUMENTO FINAL"
        score_text = f"Puntaje Cyrano: {current_score:.1f}/100 — {score_label}"
    else:
        score_text = "Puntaje Cyrano: Sin evaluar — BORRADOR"
    score_para = doc.add_paragraph(score_text)
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in score_para.runs:
        run.bold = True

    doc.add_paragraph(f"{'Idioma' if language == 'es' else 'Language'}: {'Español' if language == 'es' else 'English'}")
    doc.add_paragraph(f"{'Estado' if language == 'es' else 'Status'}: {project.status}")

    if content:
        _markdown_to_docx(doc, content)
    else:
        _generate_structured_doc(doc, project, language)

    buffer = BytesIO()
    doc.save(buffer)
    doc_bytes = buffer.getvalue()

    safe_title = project.title.replace(' ', '_')[:80]
    filename = f"{safe_title}.docx"

    existing = (
        db.query(GeneratedDoc)
        .filter(GeneratedDoc.project_id == project_id, GeneratedDoc.filename == filename)
        .first()
    )

    storage_path: str | None = None
    try:
        from app.services.storage import get_storage, make_generated_doc_key
        storage = get_storage()
        key = make_generated_doc_key(str(project_id), filename)
        storage.save(key, doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        storage_path = key
    except Exception as exc:
        import logging as _logging
        _logging.getLogger(__name__).error("generated_doc_storage_failed", extra={"error": str(exc)})

    if existing:
        existing.binary_file = None if storage_path else doc_bytes
        existing.storage_path = storage_path
        generated_doc = existing
    else:
        generated_doc = GeneratedDoc(
            project_id=project_id,
            filename=filename,
            binary_file=None if storage_path else doc_bytes,
            storage_path=storage_path,
            version_number=1,
        )
        db.add(generated_doc)
    db.commit()
    db.refresh(generated_doc)

    doc_type = "borrador" if is_draft else "documento final"
    score_msg = f" (Puntaje Cyrano: {current_score:.1f}/100)" if current_score is not None else ""

    return json.dumps({
        "status": "success",
        "document_id": str(generated_doc.id),
        "filename": filename,
        "download_url": f"/api/documents/{str(generated_doc.id)}/download",
        "is_draft": is_draft,
        "cyrano_score": current_score,
        "message": f"Documento '{filename}' generado como {doc_type}{score_msg}.",
    })


def handle_run_diagnostic(args: dict, db: Session) -> str:
    """Run the Cyrano diagnostic on a project."""
    project_id = uuid.UUID(args["project_id"]) if isinstance(args["project_id"], str) else args["project_id"]

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return json.dumps({"status": "error", "message": "Proyecto no encontrado"})

    campos_con_datos = [
        project.problem_definition,
        project.problem_tree,
        project.objectives_tree,
        project.value_chain,
        project.timeline,
        project.budget,
    ]
    tiene_full_content = bool(project.json_data and project.json_data.get("full_content"))

    if not any(campos_con_datos) and not tiene_full_content:
        return json.dumps({
            "status": "error",
            "message": "El proyecto no tiene datos guardados. Llama save_project_data antes.",
        })

    project_data = {
        "title": project.title,
        "problem_definition": project.problem_definition,
        "problem_tree": project.problem_tree,
        "objectives_tree": project.objectives_tree,
        "value_chain": project.value_chain,
        "timeline": project.timeline,
        "budget": project.budget,
    }

    jd = project.json_data or {}

    def _unwrap(val: object) -> object:
        if isinstance(val, dict) and list(val.keys()) == ["text"]:
            return val["text"]
        return val

    for struct_key, text_key in (
        ("objectives_tree", "objectives_text"),
        ("value_chain", "methodology_text"),
        ("timeline", "timeline_text"),
        ("budget", "budget_text"),
    ):
        project_data[struct_key] = _unwrap(project_data[struct_key])
        if not project_data[struct_key] and jd.get(text_key):
            project_data[struct_key] = jd[text_key]

    raw_budget = project.budget
    if isinstance(raw_budget, dict) and ("validated" in raw_budget or "issues" in raw_budget):
        project_data["budget_integrity"] = {
            "validated": raw_budget.get("validated", False),
            "total": raw_budget.get("total"),
            "issues": raw_budget.get("issues", []),
        }
    else:
        project_data["budget_integrity"] = {
            "validated": False,
            "issues": ["calculate_budget no ejecutado para este proyecto"],
        }

    if project.json_data and project.json_data.get("full_content"):
        project_data["full_content"] = project.json_data["full_content"]

    return json.dumps({
        "status": "ready_for_evaluation",
        "project_data": project_data,
        "evaluation_instructions": (
            "Evalúa cada uno de los 6 pasos con puntaje 1-10 y calcula el promedio simple: "
            "score = (suma de las 6 notas / 6) × 10. "
            "Interpreta: 1-6 Débil, 7-8 Intermedio, 9-10 Sólido. "
            "Veredicto: 'APROBADO' si score >= 95.01, 'EN REVISIÓN' si score < 95.01. "
            "OBLIGATORIO: llama save_diagnostic_result con score, sections, gaps, recommendations y verdict."
        ),
    })


def handle_save_project_data(args: dict, db: Session) -> str:
    """Save project data to PostgreSQL."""
    project_id_str = args.get("project_id")
    user_id = args.get("user_id")
    title = args.get("title") or "Proyecto sin título"

    if project_id_str:
        project = db.query(Project).filter(
            Project.id == uuid.UUID(project_id_str)
        ).first()
        if not project:
            return json.dumps({"status": "error", "message": "Proyecto no encontrado"})
        if args.get("title"):
            project.title = title
    else:
        if not user_id:
            return json.dumps({"status": "error", "message": "user_id requerido para crear proyecto"})
        project = Project(user_id=uuid.UUID(user_id), title=title)
        db.add(project)
        db.flush()
        project.title = title

    if args.get("problem_definition"):
        project.problem_definition = args["problem_definition"]

    json_data: dict = dict(project.json_data or {})
    if args.get("objectives"):
        json_data["objectives_text"] = args["objectives"]
    if args.get("methodology"):
        json_data["methodology_text"] = args["methodology"]
    if args.get("timeline_text"):
        json_data["timeline_text"] = args["timeline_text"]
    if args.get("budget_text"):
        json_data["budget_text"] = args["budget_text"]
    if args.get("full_content"):
        json_data["full_content"] = args["full_content"]
    if json_data:
        project.json_data = json_data

    project.status = ProjectStatus.in_progress
    db.commit()

    return json.dumps({
        "status": "success",
        "project_id": str(project.id),
        "message": f"Proyecto '{title}' guardado exitosamente.",
    })


def handle_save_to_project_memory(args: dict, db: Session) -> str:
    """Upload project summary to pgvector project store."""
    project_id = args["project_id"]
    summary = args["summary"]

    project_uuid = uuid.UUID(project_id) if isinstance(project_id, str) else project_id
    project = db.query(Project).filter(Project.id == project_uuid).first()
    if not project:
        return json.dumps({"status": "error", "message": "Proyecto no encontrado"})

    filename = f"proyecto_{project.title.replace(' ', '_')[:50]}.md"
    storage_id = upload_bytes_to_projects_store(
        file_bytes=summary.encode("utf-8"),
        filename=filename,
        project_id=str(project_id),
    )

    project.status = ProjectStatus.exported
    db.commit()

    return json.dumps({
        "status": "success",
        "message": f"Proyecto '{project.title}' guardado en memoria de proyectos.",
        "storage_id": storage_id,
    })


def handle_save_diagnostic_result(args: dict, db: Session) -> str:
    """Persist a structured Cyrano evaluation and update project.cyrano_score."""
    project_id = uuid.UUID(args["project_id"])

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return json.dumps({"status": "error", "message": "Proyecto no encontrado"})

    _SECTIONS = [
        "problem_definition", "problem_tree", "objectives",
        "value_chain", "timeline", "budget",
    ]
    sections = args.get("sections") or {}

    valid_notes = [
        float(sections[k])
        for k in _SECTIONS
        if k in sections and sections[k] is not None
    ]
    if valid_notes:
        calculated_score = round((sum(valid_notes) / len(valid_notes)) * 10, 2)
    else:
        calculated_score = round(float(args.get("score", 0)), 2)

    if calculated_score >= 95.01:
        verdict = "APROBADO — Sólido"
    elif calculated_score >= 70:
        verdict = "EN REVISIÓN — Intermedio"
    else:
        verdict = "EN REVISIÓN — Débil"

    existing_count = (
        db.query(CyranoEvaluation)
        .filter(CyranoEvaluation.project_id == project_id)
        .count()
    )
    version = existing_count + 1

    evaluation = CyranoEvaluation(
        project_id=project_id,
        score=calculated_score,
        sections=sections if sections else args.get("sections"),
        feedback={
            "gaps": args.get("gaps", []),
            "recommendations": args.get("recommendations", []),
        },
        verdict=verdict,
        version=version,
    )
    db.add(evaluation)
    project.cyrano_score = calculated_score
    db.commit()
    db.refresh(evaluation)

    return json.dumps({
        "status": "success",
        "evaluation_id": str(evaluation.id),
        "version": version,
        "score": calculated_score,
        "verdict": verdict,
        "message": f"Diagnóstico Cyrano v{version} persistido. Puntaje: {calculated_score:.1f}/100 — {verdict}",
    })


# ───────────────────────────────────────────
# Document generation helpers
# ───────────────────────────────────────────

def _add_tree_section(doc: Document, tree_data: dict, lang: str):
    causes_label = "Causas" if lang == "es" else "Causes"
    problem_label = "Problema Central" if lang == "es" else "Central Problem"
    effects_label = "Efectos" if lang == "es" else "Effects"

    if "central_problem" in tree_data:
        doc.add_heading(problem_label, level=2)
        doc.add_paragraph(tree_data["central_problem"])

    if "causes" in tree_data:
        doc.add_heading(causes_label, level=2)
        for i, cause in enumerate(tree_data["causes"], 1):
            doc.add_paragraph(f"{i}. {cause}", style="List Number")

    if "effects" in tree_data:
        doc.add_heading(effects_label, level=2)
        for i, effect in enumerate(tree_data["effects"], 1):
            doc.add_paragraph(f"{i}. {effect}", style="List Number")


def _add_objectives_section(doc: Document, objectives: dict, lang: str):
    general_label = "Objetivo General" if lang == "es" else "General Objective"
    specific_label = "Objetivos Específicos" if lang == "es" else "Specific Objectives"

    if "general" in objectives:
        doc.add_heading(general_label, level=2)
        doc.add_paragraph(objectives["general"])

    if "specific" in objectives:
        doc.add_heading(specific_label, level=2)
        for i, obj in enumerate(objectives["specific"], 1):
            doc.add_paragraph(f"{i}. {obj}", style="List Number")


def _add_value_chain_section(doc: Document, chain: dict, lang: str):
    if "items" in chain:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = (
            ["Actividad", "Producto", "Indicador", "Meta"] if lang == "es"
            else ["Activity", "Product", "Indicator", "Target"]
        )
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for item in chain["items"]:
            row = table.add_row()
            row.cells[0].text = item.get("actividad", item.get("activity", ""))
            row.cells[1].text = item.get("producto", item.get("product", ""))
            row.cells[2].text = item.get("indicador", item.get("indicator", ""))
            row.cells[3].text = item.get("meta", item.get("target", ""))


def _add_timeline_section(doc: Document, timeline: dict, lang: str):
    if "activities" in timeline:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = (
            ["Actividad", "Inicio", "Fin", "Responsable"] if lang == "es"
            else ["Activity", "Start", "End", "Responsible"]
        )
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for act in timeline["activities"]:
            row = table.add_row()
            row.cells[0].text = act.get("actividad", act.get("activity", ""))
            row.cells[1].text = act.get("inicio", act.get("start", ""))
            row.cells[2].text = act.get("fin", act.get("end", ""))
            row.cells[3].text = act.get("responsable", act.get("responsible", ""))


def _add_budget_section(doc: Document, budget: dict, lang: str):
    if "items" in budget:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = (
            ["Rubro", "Monto", "Actividad Vinculada"] if lang == "es"
            else ["Category", "Amount", "Linked Activity"]
        )
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for item in budget["items"]:
            row = table.add_row()
            row.cells[0].text = item.get("rubro", item.get("category", ""))
            row.cells[1].text = f"${item.get('monto', item.get('amount', 0)):,.2f}"
            row.cells[2].text = item.get("actividad_vinculada", item.get("linked_activity", ""))

        total = budget.get("total", sum(i.get("monto", i.get("amount", 0)) for i in budget["items"]))
        total_row = table.add_row()
        total_row.cells[0].text = "TOTAL"
        total_row.cells[1].text = f"${total:,.2f}"
        total_row.cells[2].text = ""


def _generate_structured_doc(doc, project, language: str):
    """Generate Word document sections from structured project data."""
    sec_titles = {
        "es": {
            "problem": "1. Identificación del Problema",
            "problem_tree": "2. Árbol de Problemas",
            "objectives": "3. Objetivos",
            "value_chain": "4. Cadena de Valor",
            "timeline": "5. Cronograma",
            "budget": "6. Presupuesto",
        },
        "en": {
            "problem": "1. Problem Identification",
            "problem_tree": "2. Problem Tree",
            "objectives": "3. Objectives",
            "value_chain": "4. Value Chain",
            "timeline": "5. Timeline",
            "budget": "6. Budget",
        }
    }
    titles = sec_titles.get(language, sec_titles["es"])

    doc.add_heading(titles["problem"], level=1)
    doc.add_paragraph(project.problem_definition or "Por definir")

    doc.add_heading(titles["problem_tree"], level=1)
    if project.problem_tree:
        _add_tree_section(doc, project.problem_tree, language)
    else:
        doc.add_paragraph("Por definir")

    doc.add_heading(titles["objectives"], level=1)
    if project.objectives_tree:
        _add_objectives_section(doc, project.objectives_tree, language)
    else:
        doc.add_paragraph("Por definir")

    doc.add_heading(titles["value_chain"], level=1)
    if project.value_chain:
        _add_value_chain_section(doc, project.value_chain, language)
    else:
        doc.add_paragraph("Por definir")

    doc.add_heading(titles["timeline"], level=1)
    if project.timeline:
        _add_timeline_section(doc, project.timeline, language)
    else:
        doc.add_paragraph("Por definir")

    doc.add_heading(titles["budget"], level=1)
    if project.budget:
        _add_budget_section(doc, project.budget, language)
    else:
        doc.add_paragraph("Por definir")


def _markdown_to_docx(doc, markdown_text: str):
    """Convert markdown text to Word document content."""
    lines = markdown_text.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=level)
            continue

        if stripped.startswith(('- ', '• ', '* ')):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            continue

        num_match = re.match(r'^(\d+)[\.\)]\s+(.+)$', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, num_match.group(2))
            continue

        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)


def _add_formatted_text(paragraph, text: str):
    """Add text with bold formatting to a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
