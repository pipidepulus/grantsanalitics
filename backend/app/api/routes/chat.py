import uuid
import json
import io
import logging
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from slowapi import Limiter
from slowapi.util import get_remote_address
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_db
from app.models.conversation import Conversation
from app.schemas.chat import ChatRequest, ChatResponse
from app.services.ai_agent import process_chat_message, process_chat_message_stream

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/chat", tags=["chat"])

# Re-use the limiter singleton from main; instantiating a new one is harmless
# because slowapi reads the rate from the decorator, not the instance.
_limiter = Limiter(key_func=get_remote_address)


@router.post("", response_model=ChatResponse)
@_limiter.limit("30/minute")
async def chat(request: Request, chat_req: ChatRequest):
    """Send a message to the AI agent and get a response."""
    try:
        result = await process_chat_message(
            user_id=chat_req.user_id,
            message=chat_req.message,
            conversation_id=chat_req.conversation_id,
            project_id=chat_req.project_id,
        )
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.exception("Error processing chat message")
        raise HTTPException(status_code=500, detail="Error al procesar tu mensaje. Por favor, inténtalo de nuevo.")
    return result


@router.post("/stream")
@_limiter.limit("30/minute")
async def chat_stream(request: Request, chat_req: ChatRequest):
    """Stream a chat response using Server-Sent Events."""
    async def event_generator():
        try:
            async for event in process_chat_message_stream(
                user_id=chat_req.user_id,
                message=chat_req.message,
                conversation_id=chat_req.conversation_id,
                project_id=chat_req.project_id,
            ):
                yield f"data: {json.dumps(event)}\n\n"
        except ValueError as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"
        except Exception as e:
            logger.exception("Error in streaming chat")
            yield f"data: {json.dumps({'type': 'error', 'content': 'Error al procesar tu mensaje.'})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/conversations/{user_id}")
def list_conversations(
    user_id: uuid.UUID,
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    db: Session = Depends(get_db),
):
    """List all conversations for a user."""
    conversations = (
        db.query(Conversation)
        .filter(Conversation.user_id == user_id)
        .order_by(Conversation.updated_at.desc())
        .offset(skip)
        .limit(limit)
        .all()
    )
    return [
        {
            "id": c.id,
            "title": c.title,
            "project_id": c.project_id,
            "created_at": c.created_at,
            "updated_at": c.updated_at,
            "message_count": len(c.messages),
        }
        for c in conversations
    ]


@router.delete("/conversations/{user_id}/{conversation_id}")
def delete_conversation(user_id: uuid.UUID, conversation_id: uuid.UUID, db: Session = Depends(get_db)):
    """Delete a conversation and all its messages."""
    conv = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == user_id)
        .first()
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    from app.models.conversation import Message
    db.query(Message).filter(Message.conversation_id == conv.id).delete()
    db.delete(conv)
    db.commit()
    return {"ok": True}


@router.get("/conversations/{user_id}/{conversation_id}")
def get_conversation(user_id: uuid.UUID, conversation_id: uuid.UUID, db: Session = Depends(get_db)):
    """Get a full conversation with messages."""
    conv = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == user_id)
        .first()
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    return {
        "id": conv.id,
        "title": conv.title,
        "project_id": conv.project_id,
        "messages": [
            {
                "id": m.id,
                "role": m.role,
                "content": m.content,
                "tool_calls": m.tool_calls,
                "created_at": m.created_at,
            }
            for m in conv.messages
        ],
    }


@router.post("/conversations/{user_id}/{conversation_id}/save")
def save_conversation(user_id: uuid.UUID, conversation_id: uuid.UUID, project_id: uuid.UUID | None = None, db: Session = Depends(get_db)):
    """Save a conversation as a Word document in the database."""
    conv = (
        db.query(Conversation)
        .filter(Conversation.id == conversation_id, Conversation.user_id == user_id)
        .first()
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    # Use project_id from query param, or fall back to conversation's project_id
    effective_project_id = project_id or conv.project_id
    if not effective_project_id:
        raise HTTPException(status_code=400, detail="No hay proyecto vinculado. Selecciona un proyecto en la barra lateral primero.")

    # Sync project_id to conversation if it wasn't set
    if not conv.project_id and effective_project_id:
        conv.project_id = effective_project_id
        db.commit()

    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(conv.title or "Conversación", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Guardado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    for msg in conv.messages:
        # Role header
        role_label = "👤 Usuario" if msg.role == "user" else "🤖 Pipidepulus AI"
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(role_label)
        role_run.bold = True
        role_run.font.size = Pt(11)
        if msg.role == "assistant":
            role_run.font.color.rgb = RGBColor(59, 130, 246)

        # Timestamp
        if msg.created_at:
            ts = msg.created_at
            if isinstance(ts, str):
                ts = datetime.fromisoformat(ts)
            time_run = role_para.add_run(f"  — {ts.strftime('%d/%m/%Y %H:%M')}")
            time_run.font.size = Pt(8)
            time_run.font.color.rgb = RGBColor(160, 160, 160)

        # Content
        content = msg.content or ""
        for line in content.split("\n"):
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(10)

        # Separator
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 60)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)

    buf = io.BytesIO()
    doc.save(buf)
    binary_data = buf.getvalue()

    # Use project title for filename, not conversation title
    from app.models.project import Project
    project = db.query(Project).filter(Project.id == effective_project_id).first()
    project_title = project.title if project else (conv.title or "conversacion")
    safe_title = project_title.replace(" ", "_")[:80]
    filename = f"Chat_{safe_title}.docx"

    # Update existing document or create new one
    from app.models.document import GeneratedDoc
    existing = (
        db.query(GeneratedDoc)
        .filter(GeneratedDoc.project_id == effective_project_id, GeneratedDoc.filename == filename)
        .first()
    )

    # Write to object storage; fall back to inline binary on failure
    storage_path: str | None = None
    try:
        from app.services.storage import get_storage, make_generated_doc_key
        _storage = get_storage()
        key = make_generated_doc_key(str(effective_project_id), filename)
        _storage.save(key, binary_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        storage_path = key
    except Exception as _exc:
        logger.error("chat_doc_storage_failed", extra={"error": str(_exc)})

    if existing:
        existing.binary_file = None if storage_path else binary_data
        existing.storage_path = storage_path
        gen_doc = existing
    else:
        gen_doc = GeneratedDoc(
            project_id=effective_project_id,
            filename=filename,
            binary_file=None if storage_path else binary_data,
            storage_path=storage_path,
            version_number=1,
        )
        db.add(gen_doc)

    db.commit()
    db.refresh(gen_doc)

    return {
        "id": str(gen_doc.id),
        "filename": gen_doc.filename,
        "message": f"Conversación guardada como '{filename}'",
    }
